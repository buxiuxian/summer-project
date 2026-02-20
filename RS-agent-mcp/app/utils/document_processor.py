"""
文档处理工具 - 用于提取PDF和Word文档的文本内容
"""

import logging
import subprocess
import tempfile
import os
from pathlib import Path
from typing import Optional

# PDF处理库
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word文档处理库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# DOC文件处理库 - 使用多种方法
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

# CSV和XLSX处理库
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

logger = logging.getLogger(__name__)

def extract_text_from_csv(file_path: str) -> Optional[str]:
    """
    从CSV文件中提取文本内容，保留表格结构
    
    Args:
        file_path: CSV文件路径
        
    Returns:
        提取的文本内容，失败时返回None
    """
    if not PANDAS_AVAILABLE:
        logger.error("pandas库未安装，无法处理CSV文件")
        return None
    
    try:
        # 尝试不同的编码格式读取CSV
        encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'latin-1']
        df = None
        
        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                logger.info(f"成功使用{encoding}编码读取CSV文件")
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"使用{encoding}编码读取CSV失败: {str(e)}")
                continue
        
        if df is None:
            logger.error("所有编码格式都无法读取CSV文件")
            return None
        
        # 生成表格结构化文本
        text_content = f"CSV文件内容（共{len(df)}行，{len(df.columns)}列）：\n\n"
        
        # 添加列标题
        text_content += "列标题：\n"
        text_content += " | ".join([str(col) for col in df.columns]) + "\n"
        text_content += "-" * (len(" | ".join([str(col) for col in df.columns]))) + "\n"
        
        # 添加数据行（限制显示前100行避免内容过长）
        max_rows = min(100, len(df))
        for index, row in df.head(max_rows).iterrows():
            row_text = " | ".join([str(value) if not pd.isna(value) else "" for value in row])
            text_content += row_text + "\n"
        
        if len(df) > max_rows:
            text_content += f"\n... 还有{len(df) - max_rows}行数据未显示\n"
        
        # 添加统计信息
        text_content += f"\n数据统计信息：\n"
        text_content += f"总行数：{len(df)}\n"
        text_content += f"总列数：{len(df.columns)}\n"
        
        # 对数值列添加基本统计
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            text_content += f"\n数值列统计：\n"
            for col in numeric_cols:
                text_content += f"{col}: 平均值={df[col].mean():.2f}, 最小值={df[col].min()}, 最大值={df[col].max()}\n"
        
        logger.info(f"成功从CSV提取文本，长度: {len(text_content)} 字符")
        return text_content
        
    except Exception as e:
        logger.error(f"CSV文本提取出错: {str(e)}")
        return None

def extract_text_from_xlsx(file_path: str) -> Optional[str]:
    """
    从XLSX文件中提取文本内容，保留表格结构
    
    Args:
        file_path: XLSX文件路径
        
    Returns:
        提取的文本内容，失败时返回None
    """
    if not PANDAS_AVAILABLE:
        logger.error("pandas库未安装，无法处理XLSX文件")
        return None
    
    try:
        # 读取Excel文件（默认读取第一个工作表）
        df = pd.read_excel(file_path, engine='openpyxl')
        
        # 生成表格结构化文本
        text_content = f"XLSX文件内容（共{len(df)}行，{len(df.columns)}列）：\n\n"
        
        # 添加列标题
        text_content += "列标题：\n"
        text_content += " | ".join([str(col) for col in df.columns]) + "\n"
        text_content += "-" * (len(" | ".join([str(col) for col in df.columns]))) + "\n"
        
        # 添加数据行（限制显示前100行避免内容过长）
        max_rows = min(100, len(df))
        for index, row in df.head(max_rows).iterrows():
            row_text = " | ".join([str(value) if not pd.isna(value) else "" for value in row])
            text_content += row_text + "\n"
        
        if len(df) > max_rows:
            text_content += f"\n... 还有{len(df) - max_rows}行数据未显示\n"
        
        # 添加统计信息
        text_content += f"\n数据统计信息：\n"
        text_content += f"总行数：{len(df)}\n"
        text_content += f"总列数：{len(df.columns)}\n"
        
        # 对数值列添加基本统计
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            text_content += f"\n数值列统计：\n"
            for col in numeric_cols:
                text_content += f"{col}: 平均值={df[col].mean():.2f}, 最小值={df[col].min()}, 最大值={df[col].max()}\n"
        
        # 如果Excel文件有多个工作表，提示用户
        try:
            xl_file = pd.ExcelFile(file_path)
            if len(xl_file.sheet_names) > 1:
                text_content += f"\n注意：此Excel文件包含{len(xl_file.sheet_names)}个工作表：{', '.join(xl_file.sheet_names)}\n"
                text_content += "当前只显示了第一个工作表的内容。\n"
        except:
            pass
        
        logger.info(f"成功从XLSX提取文本，长度: {len(text_content)} 字符")
        return text_content
        
    except Exception as e:
        logger.error(f"XLSX文本提取出错: {str(e)}")
        return None

def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    从PDF文件中提取文本内容
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        提取的文本内容，失败时返回None
    """
    if not PDF_AVAILABLE:
        logger.error("PDF处理库未安装，无法处理PDF文件")
        return None
    
    try:
        # 优先使用pdfplumber，它对复杂布局的处理更好
        text_content = _extract_with_pdfplumber(file_path)
        
        if not text_content or len(text_content.strip()) < 50:
            # 如果pdfplumber提取的内容太少，尝试PyPDF2
            logger.info("pdfplumber提取内容较少，尝试使用PyPDF2")
            text_content = _extract_with_pypdf2(file_path)
        
        if text_content and len(text_content.strip()) > 10:
            logger.info(f"成功从PDF提取文本，长度: {len(text_content)} 字符")
            return text_content
        else:
            logger.warning("PDF文本提取失败或内容为空")
            return None
            
    except Exception as e:
        logger.error(f"PDF文本提取出错: {str(e)}")
        return None

def _extract_with_pdfplumber(file_path: str) -> Optional[str]:
    """使用pdfplumber提取PDF文本"""
    try:
        with pdfplumber.open(file_path) as pdf:
            text_parts = []
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- 第{page_num}页 ---\n{page_text}\n")
            
            return '\n'.join(text_parts) if text_parts else None
            
    except Exception as e:
        logger.error(f"pdfplumber提取失败: {str(e)}")
        return None

def _extract_with_pypdf2(file_path: str) -> Optional[str]:
    """使用PyPDF2提取PDF文本"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- 第{page_num}页 ---\n{page_text}\n")
            
            return '\n'.join(text_parts) if text_parts else None
            
    except Exception as e:
        logger.error(f"PyPDF2提取失败: {str(e)}")
        return None

def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    从DOCX文件中提取文本内容
    
    Args:
        file_path: DOCX文件路径
        
    Returns:
        提取的文本内容，失败时返回None
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx库未安装，无法处理DOCX文件")
        return None
    
    try:
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 处理表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        text_content = '\n'.join(text_parts)
        
        if text_content and len(text_content.strip()) > 10:
            logger.info(f"成功从DOCX提取文本，长度: {len(text_content)} 字符")
            return text_content
        else:
            logger.warning("DOCX文本提取失败或内容为空")
            return None
            
    except Exception as e:
        logger.error(f"DOCX文本提取出错: {str(e)}")
        return None

def _extract_doc_with_antiword(file_path: str) -> Optional[str]:
    """
    使用antiword命令行工具提取DOC文件文本
    antiword是专门处理传统DOC格式的工具
    """
    try:
        # 尝试调用antiword命令
        result = subprocess.run(
            ['antiword', file_path],
            capture_output=True,
            text=True,
            timeout=30,
            encoding='utf-8'
        )
        
        if result.returncode == 0 and result.stdout:
            text_content = result.stdout.strip()
            if len(text_content) > 10:
                logger.info(f"antiword成功提取DOC文本，长度: {len(text_content)} 字符")
                return text_content
        
        logger.warning(f"antiword提取失败，返回码: {result.returncode}")
        if result.stderr:
            logger.warning(f"antiword错误信息: {result.stderr}")
        return None
        
    except FileNotFoundError:
        logger.warning("antiword命令未找到，请安装antiword工具")
        return None
    except subprocess.TimeoutExpired:
        logger.error("antiword处理超时")
        return None
    except Exception as e:
        logger.error(f"antiword调用出错: {str(e)}")
        return None

def _extract_doc_with_libreoffice(file_path: str) -> Optional[str]:
    """
    使用LibreOffice命令行工具转换DOC文件
    """
    try:
        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            # 使用LibreOffice转换为文本文件
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'txt:Text',
                '--outdir', temp_dir, file_path
            ], capture_output=True, timeout=60)
            
            if result.returncode == 0:
                # 查找生成的txt文件
                doc_name = Path(file_path).stem
                txt_file = Path(temp_dir) / f"{doc_name}.txt"
                
                if txt_file.exists():
                    with open(txt_file, 'r', encoding='utf-8') as f:
                        text_content = f.read().strip()
                    
                    if len(text_content) > 10:
                        logger.info(f"LibreOffice成功转换DOC文件，长度: {len(text_content)} 字符")
                        return text_content
            
            logger.warning("LibreOffice转换失败")
            return None
            
    except FileNotFoundError:
        logger.warning("LibreOffice命令未找到")
        return None
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice转换超时")
        return None
    except Exception as e:
        logger.error(f"LibreOffice转换出错: {str(e)}")
        return None

def _extract_doc_with_docx2txt_fallback(file_path: str) -> Optional[str]:
    """
    使用docx2txt作为降级处理（某些DOC文件可能兼容）
    """
    if not DOCX2TXT_AVAILABLE:
        return None
    
    try:
        # 某些DOC文件可能与DOCX兼容
        text_content = docx2txt.process(file_path)
        
        if text_content and len(text_content.strip()) > 10:
            logger.info(f"docx2txt降级处理成功，长度: {len(text_content)} 字符")
            return text_content
        
        return None
        
    except Exception as e:
        logger.debug(f"docx2txt降级处理失败: {str(e)}")
        return None

def extract_text_from_doc(file_path: str) -> Optional[str]:
    """
    从DOC文件中提取文本内容
    使用多种方法尝试处理传统DOC格式文件
    
    Args:
        file_path: DOC文件路径
        
    Returns:
        提取的文本内容，失败时返回None
    """
    logger.info(f"开始处理DOC文件: {file_path}")
    
    # 方法1: 使用antiword（最可靠）
    text_content = _extract_doc_with_antiword(file_path)
    if text_content:
        return text_content
    
    # 方法2: 使用LibreOffice（通用性好）
    text_content = _extract_doc_with_libreoffice(file_path)
    if text_content:
        return text_content
    
    # 方法3: 尝试docx2txt降级处理
    text_content = _extract_doc_with_docx2txt_fallback(file_path)
    if text_content:
        return text_content
    
    # 方法4: 尝试python-docx（最后的尝试）
    try:
        text_content = extract_text_from_docx(file_path)
        if text_content:
            logger.info("python-docx兼容性处理成功")
            return text_content
    except Exception as e:
        logger.debug(f"python-docx兼容性处理失败: {str(e)}")
    
    logger.error(f"所有DOC文件处理方法都失败了: {file_path}")
    return None

def extract_document_text(file_path: str) -> Optional[str]:
    """
    根据文件类型自动选择合适的方法提取文档文本
    
    Args:
        file_path: 文档文件路径
        
    Returns:
        提取的文本内容，失败时返回None
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        logger.error(f"文件不存在: {file_path}")
        return None
    
    extension = file_path.suffix.lower()
    
    if extension == '.pdf':
        return extract_text_from_pdf(str(file_path))
    elif extension == '.docx':
        return extract_text_from_docx(str(file_path))
    elif extension == '.csv':
        return extract_text_from_csv(str(file_path))
    elif extension == '.xlsx':
        return extract_text_from_xlsx(str(file_path))
    else:
        logger.error(f"不支持的文档类型: {extension}")
        return None

def validate_extracted_text(text: str) -> bool:
    """
    验证提取的文本是否有效
    
    Args:
        text: 提取的文本
        
    Returns:
        是否有效
    """
    if not text or not isinstance(text, str):
        return False
    
    # 去除空白字符后的长度
    clean_text = text.strip()
    
    # 文本长度至少20个字符
    if len(clean_text) < 20:
        return False
    
    # 检查是否包含足够的字母字符（避免纯符号或数字）
    letter_count = sum(1 for c in clean_text if c.isalpha())
    if letter_count < len(clean_text) * 0.3:  # 至少30%是字母
        return False
    
    return True

def check_doc_processing_tools() -> dict:
    """
    检查DOC文件处理工具的可用性
    
    Returns:
        各种工具的可用性状态
    """
    tools_status = {
        'antiword': False,
        'libreoffice': False,
        'docx2txt': DOCX2TXT_AVAILABLE,
        'python_docx': DOCX_AVAILABLE
    }
    
    # 检查antiword
    try:
        result = subprocess.run(['antiword', '-v'], capture_output=True, timeout=5)
        tools_status['antiword'] = result.returncode == 0
    except:
        pass
    
    # 检查LibreOffice
    try:
        result = subprocess.run(['libreoffice', '--version'], capture_output=True, timeout=5)
        tools_status['libreoffice'] = result.returncode == 0
    except:
        pass
    
    return tools_status 